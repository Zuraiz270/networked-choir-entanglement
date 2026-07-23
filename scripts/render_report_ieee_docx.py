"""Create an editable Word counterpart of the IEEE-style report.

The document uses an explicit IEEE-style fidelity override: A4 paper, Times
typography, a full-width title and abstract, then two local reading columns for
the report body. This preserves the layout intent of the PDF while keeping the
result editable in Word.

Usage:
    python scripts/render_report_ieee_docx.py
"""

from __future__ import annotations

import argparse
import html
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor, Twips
from scripts.render_report_ieee import (
    _IMAGE_RE,
    _TABLE_TITLES,
    DEFAULT_SOURCE,
    _front_matter,
    _is_special,
    _main_heading,
    _roman,
    _subheading,
)

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_DOCX_OUTPUT = ROOT / "output/docx/networked_choir_final_report_ieee.docx"

FONT_NAME = "Times New Roman"
CODE_FONT_NAME = "Courier New"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(70, 70, 70)
BODY_SIZE = 9.2
BODY_LEADING = 10.8
COLUMN_IMAGE_WIDTH = Cm(8.0)
TABLE_WIDTH_DXA = int(Cm(8.0).twips)
TABLE_INDENT_DXA = 120
TABLE_CELL_MARGINS_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
COLUMN_GAP_DXA = 260

_INLINE_TOKEN_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|\[[^]]+\]\([^)]+\))")
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _set_run_font(
    run: Any,
    *,
    name: str = FONT_NAME,
    size: float | None = None,
    color: RGBColor = BLACK,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _configure_paragraph_format(
    paragraph_format: Any,
    *,
    before: float = 0,
    after: float = 0,
    line: float | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    first_indent: float | None = None,
    keep_next: bool = False,
) -> None:
    paragraph_format.space_before = Pt(before)
    paragraph_format.space_after = Pt(after)
    if line is not None:
        paragraph_format.line_spacing = Pt(line)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    if alignment is not None:
        paragraph_format.alignment = alignment
    if first_indent is not None:
        paragraph_format.first_line_indent = Cm(first_indent)
    paragraph_format.keep_with_next = keep_next


def _configure_style(
    style: Any,
    *,
    size: float,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor = BLACK,
    before: float = 0,
    after: float = 0,
    line: float | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
    first_indent: float | None = None,
    keep_next: bool = False,
) -> None:
    style.font.name = FONT_NAME
    style._element.rPr.rFonts.set(qn("w:ascii"), FONT_NAME)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_NAME)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = color
    _configure_paragraph_format(
        style.paragraph_format,
        before=before,
        after=after,
        line=line,
        alignment=alignment,
        first_indent=first_indent,
        keep_next=keep_next,
    )


def _configure_styles(doc: Document) -> None:
    styles = doc.styles
    _configure_style(
        styles["Normal"],
        size=BODY_SIZE,
        after=2.2,
        line=BODY_LEADING,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_indent=0.35,
    )
    for name in ("List Bullet", "List Number"):
        _configure_style(styles[name], size=BODY_SIZE, after=1.5, line=BODY_LEADING)

    custom_styles = {
        "IEEE Title": (21, True, False, BLACK, 0, 9, 23, WD_ALIGN_PARAGRAPH.CENTER, None, True),
        "IEEE Authors": (11, False, False, BLACK, 0, 4, 13, WD_ALIGN_PARAGRAPH.CENTER, None, False),
        "IEEE Affiliation": (
            8.5,
            False,
            False,
            BLACK,
            0,
            2,
            10.5,
            WD_ALIGN_PARAGRAPH.CENTER,
            None,
            False,
        ),
        "IEEE Abstract": (
            8.5,
            False,
            False,
            BLACK,
            0,
            4,
            10.2,
            WD_ALIGN_PARAGRAPH.JUSTIFY,
            None,
            False,
        ),
        "IEEE Index": (
            8.5,
            False,
            False,
            BLACK,
            0,
            9,
            10.2,
            WD_ALIGN_PARAGRAPH.JUSTIFY,
            None,
            False,
        ),
        "IEEE Caption": (8, False, False, MUTED, 3, 7, 9.4, WD_ALIGN_PARAGRAPH.CENTER, None, False),
        "IEEE Table Caption": (
            7.8,
            False,
            False,
            BLACK,
            5,
            3,
            9,
            WD_ALIGN_PARAGRAPH.CENTER,
            None,
            True,
        ),
        "IEEE Table Cell": (
            7.1,
            False,
            False,
            BLACK,
            0,
            0,
            8.3,
            WD_ALIGN_PARAGRAPH.LEFT,
            None,
            False,
        ),
    }
    for name, values in custom_styles.items():
        style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        _configure_style(
            style,
            size=values[0],
            bold=values[1],
            italic=values[2],
            color=values[3],
            before=values[4],
            after=values[5],
            line=values[6],
            alignment=values[7],
            first_indent=values[8],
            keep_next=values[9],
        )

    heading_1 = styles["Heading 1"]
    _configure_style(
        heading_1,
        size=10,
        after=4,
        line=12,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        keep_next=True,
    )
    heading_2 = styles["Heading 2"]
    _configure_style(
        heading_2,
        size=9.5,
        italic=True,
        after=3,
        line=11.3,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        keep_next=True,
    )


def _set_columns(section: Any, count: int, gap_dxa: int = COLUMN_GAP_DXA) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(gap_dxa))


def _configure_section(section: Any, *, columns: int) -> None:
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.header_distance = Cm(0.5)
    section.footer_distance = Cm(0.5)
    _set_columns(section, columns)


def _add_page_field(paragraph: Any) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    _set_run_font(run, size=7.5, color=MUTED)


def _configure_footer(section: Any) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _configure_paragraph_format(paragraph.paragraph_format, after=0, line=9)
    _add_page_field(paragraph)


def _add_inline_runs(paragraph: Any, text: str, *, size: float = BODY_SIZE) -> None:
    escaped_text = html.unescape(text)
    cursor = 0
    for match in _INLINE_TOKEN_RE.finditer(escaped_text):
        if match.start() > cursor:
            run = paragraph.add_run(escaped_text[cursor : match.start()])
            _set_run_font(run, size=size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, size=size, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, name=CODE_FONT_NAME, size=size)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, size=size, italic=True)
        else:
            label = re.match(r"\[([^]]+)\]", token)
            run = paragraph.add_run(label.group(1) if label else token)
            _set_run_font(run, size=size, color=RGBColor(0, 0, 128))
            run.underline = True
        cursor = match.end()
    if cursor < len(escaped_text):
        run = paragraph.add_run(escaped_text[cursor:])
        _set_run_font(run, size=size)


def _add_paragraph(doc: Document, text: str, style: str = "Normal") -> Any:
    paragraph = doc.add_paragraph(style=style)
    _add_inline_runs(paragraph, text, size=BODY_SIZE if style == "Normal" else None or BODY_SIZE)
    return paragraph


def _add_labeled_paragraph(doc: Document, label: str, value: str, style: str) -> Any:
    paragraph = doc.add_paragraph(style=style)
    label_run = paragraph.add_run(label)
    _set_run_font(label_run, size=8.5, bold=True)
    _add_inline_runs(paragraph, value, size=8.5)
    return paragraph


def _set_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "bottom", "insideH"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
    for edge in ("left", "right", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def _set_cell_margins(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in TABLE_CELL_MARGINS_DXA.items():
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _apply_table_geometry(table: Any, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        element = OxmlElement("w:gridCol")
        element.set(qn("w:w"), str(width))
        grid.append(element)
    for col, width in enumerate(widths):
        table.columns[col].width = Twips(width)
    for row in table.rows:
        row.height = None
        for col, cell in enumerate(row.cells):
            cell.width = Twips(widths[col])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[col]))
            _set_cell_margins(cell)
    _set_table_borders(table)


def _add_table(doc: Document, lines: list[str], table_number: int, styles: dict[str, Any]) -> None:
    rows = [[cell.strip() for cell in line.strip().strip("|").split("|")] for line in lines]
    if len(rows) > 1 and all(set(cell) <= {":", "-"} for cell in rows[1]):
        rows.pop(1)
    column_count = max(len(row) for row in rows)
    for row in rows:
        row.extend("" for _ in range(column_count - len(row)))
    table = doc.add_table(rows=len(rows), cols=column_count)
    widths = [TABLE_WIDTH_DXA // column_count] * column_count
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    _apply_table_geometry(table, widths)
    for row_idx, row in enumerate(rows):
        for col_idx, text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            paragraph = cell.paragraphs[0]
            paragraph.style = styles["IEEE Table Cell"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_inline_runs(paragraph, text, size=7.1)
            if row_idx == 0:
                for run in paragraph.runs:
                    run.bold = True
    header_row = table.rows[0]._tr
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    header_row.get_or_add_trPr().append(header)
    caption = doc.add_paragraph(style="IEEE Table Caption")
    caption.add_run(f"TABLE {_roman(table_number)}\n{_TABLE_TITLES[table_number - 1].upper()}")
    for run in caption.runs:
        _set_run_font(run, size=7.8)
    caption._p.addnext(table._tbl)
    spacer = doc.add_paragraph()
    _configure_paragraph_format(spacer.paragraph_format, after=2, line=3)


def _add_image(doc: Document, source: Path, line: str) -> None:
    match = _IMAGE_RE.match(line.strip())
    if match is None:
        raise ValueError(f"invalid report image: {line}")
    image_path = (source.parent / match.group("path")).resolve()
    if not image_path.exists():
        raise FileNotFoundError(f"report image not found: {image_path}")
    image_paragraph = doc.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True
    run = image_paragraph.add_run()
    run.add_picture(str(image_path), width=COLUMN_IMAGE_WIDTH)
    caption_text = re.sub(r"^Figure\s+(\d+)\.", r"Fig. \1.", match.group("caption"))
    caption = doc.add_paragraph(style="IEEE Caption")
    _add_inline_runs(caption, caption_text, size=8)


def _add_list(doc: Document, items: list[str], ordered: bool) -> None:
    style = "List Number" if ordered else "List Bullet"
    for item in items:
        paragraph = doc.add_paragraph(style=style)
        _add_inline_runs(paragraph, item, size=BODY_SIZE)


def _add_body(doc: Document, source: Path, lines: list[str], styles: dict[str, Any]) -> None:
    index = 0
    table_number = 0
    subsection_index = 0
    while index < len(lines):
        line = lines[index].strip()
        if not line or line in {"---", "<!-- pagebreak -->"}:
            index += 1
            continue
        if line.startswith("## "):
            subsection_index = 0
            paragraph = doc.add_paragraph(style="Heading 1")
            paragraph.add_run(_main_heading(line[3:]))
            for run in paragraph.runs:
                _set_run_font(run, size=10)
            index += 1
            continue
        if line.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            paragraph.add_run(_subheading(line[4:], subsection_index))
            for run in paragraph.runs:
                _set_run_font(run, size=9.5, italic=True)
            subsection_index += 1
            index += 1
            continue
        if _IMAGE_RE.match(line):
            _add_image(doc, source, line)
            index += 1
            continue
        if line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            table_number += 1
            _add_table(doc, table_lines, table_number, styles)
            continue
        if line.startswith("- "):
            items: list[str] = []
            while index < len(lines) and lines[index].strip().startswith("- "):
                items.append(lines[index].strip()[2:])
                index += 1
            _add_list(doc, items, ordered=False)
            continue
        if re.match(r"^\d+\.\s", line):
            items = []
            while index < len(lines) and re.match(r"^\d+\.\s", lines[index].strip()):
                items.append(re.sub(r"^\d+\.\s+", "", lines[index].strip()))
                index += 1
            _add_list(doc, items, ordered=True)
            continue
        paragraph_lines = [line]
        index += 1
        while index < len(lines) and not _is_special(lines[index]):
            paragraph_lines.append(lines[index].strip())
            index += 1
        _add_paragraph(doc, " ".join(paragraph_lines))


def build_docx(source: Path, output: Path) -> None:
    """Build an editable Word document from the canonical Markdown report."""
    source = source.resolve()
    output = output.resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    lines = source.read_text(encoding="utf-8").splitlines()
    title, metadata, abstract, body_lines = _front_matter(lines)

    doc = Document()
    _configure_styles(doc)
    first_section = doc.sections[0]
    _configure_section(first_section, columns=1)
    _configure_footer(first_section)

    title_paragraph = doc.add_paragraph(style="IEEE Title")
    _add_inline_runs(title_paragraph, title, size=21)
    authors = doc.add_paragraph(style="IEEE Authors")
    _add_inline_runs(authors, metadata["Authors"], size=11)
    for label in ("Course", "Supervisors", "Date"):
        affiliation = doc.add_paragraph(style="IEEE Affiliation")
        _add_inline_runs(affiliation, metadata[label], size=8.5)
    spacer = doc.add_paragraph()
    _configure_paragraph_format(spacer.paragraph_format, after=4, line=6)
    _add_labeled_paragraph(doc, "Abstract- ", abstract, "IEEE Abstract")
    _add_labeled_paragraph(doc, "Index Terms- ", metadata["Keywords"], "IEEE Index")

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    _configure_section(body_section, columns=2)
    _configure_footer(body_section)
    _add_body(doc, source, body_lines, {name: doc.styles[name] for name in ("IEEE Table Cell",)})

    properties = doc.core_properties
    properties.title = title
    properties.author = metadata["Authors"]
    properties.subject = "IEEE-style seminar report"
    properties.keywords = metadata["Keywords"]
    doc.save(output)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source", nargs="?", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("output", nargs="?", type=Path, default=DEFAULT_DOCX_OUTPUT)
    args = parser.parse_args()
    build_docx(args.source, args.output)
    print(f"Wrote {args.output}")


if __name__ == "__main__":
    main()
